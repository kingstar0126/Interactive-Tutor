from flask import Blueprint, jsonify, request, current_app
from . import db
from rich import print, pretty
import datetime
from bs4 import BeautifulSoup
import requests
import os
import json
from .models import Chat, Train, User, Invite
import re
from langchain.docstore.document import Document
from typing import List
from pypdf import PdfReader
from .traindata import *
from io import BytesIO
import pandas as pd
import docx
from ebooklib import epub
import ebooklib
import pinecone
import openai
from dotenv import load_dotenv
from werkzeug.exceptions import RequestEntityTooLarge
from werkzeug.utils import secure_filename
from werkzeug.datastructures import FileStorage
import csv
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.service import Service as ChromeService
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from webdriver_manager.chrome import ChromeDriverManager
from urllib.parse import urljoin
from urllib.parse import urlparse
import boto3
from threading import Thread

load_dotenv()

PINECONE_API_KEY = os.getenv('PINECONE_API_KEY')
PINECONE_ENV = os.getenv('PINECONE_ENV')
OPENAI_API_KEY = os.getenv('OPENAI_API_KEY')
PINECONE_INDEX_NAME = os.getenv('PINECONE_INDEX_NAME')
pinecone.init(api_key=PINECONE_API_KEY, environment=PINECONE_ENV)
openai.openai_api_key = OPENAI_API_KEY

pretty.install()

train = Blueprint('train', __name__)

S3_CLIENT = boto3.client(
    's3',
    aws_access_key_id=os.getenv('ACCESS_KEY_ID'),
    aws_secret_access_key=os.getenv('ACCESS_SECRET_KEY'),
    region_name=os.getenv('REGION')
)
S3_PRIVATE_BUCKET = os.getenv('S3_PRIVATE')

def compare_token_words(ct, chatbot):
    current_chat = db.session.query(Chat).filter_by(uuid=chatbot).first()
    user = db.session.query(User).filter_by(id=current_chat.user_id).first()

    # if user.role == 1 or user.role == 7:
    #     return True

    # elif ct >= user.training_words:
    #     return False
    return True


def delete_vectore(source, chat):
    try:
        index = pinecone.Index(PINECONE_INDEX_NAME)
        return index.delete(
            filter={
                "source": f"{source}",
                "chat": f"{chat}"
            }
        )
    except Exception as e:
        print(f"delete_vectore error: {e}")


def create_vector(docs):
    embeddings = OpenAIEmbeddings(openai_api_key=OPENAI_API_KEY)
    Pinecone.from_documents(
        docs, embeddings, index_name=PINECONE_INDEX_NAME)


def parse_pdf(file: BytesIO) -> List[str]:
    pdf = PdfReader(file)
    output = []
    for page in pdf.pages:
        text = page.extract_text()
        # print(text)
        # # Merge hyphenated words
        # text = re.sub(r"(\w+)-\n(\w+)", r"\1\2", text)
        # # Fix newlines in the middle of sentences
        # text = re.sub(r"(?<!\n\s)\n(?!\s\n)", " ", text.strip())
        # # Remove multiple newlines
        # text = re.sub(r"\n\s*\n", "\n\n", text)
        # text = correct_grammar(text)
        output.append(text)
    return output


def parse_csv(file):
    df = pd.read_csv(file)
    # Convert DataFrame to list of dictionaries and then each dictionary to string
    return [str(record) for record in df.to_dict('records')]

def extract_data_from_xlsx(file):
    df = pd.read_excel(file)
    # Convert DataFrame to list of dictionaries and then each dictionary to string
    return [str(record) for record in df.to_dict('records')]

def parse_docx(file):
    doc = docx.Document(file)
    fullText = []
    for para in doc.paragraphs:
        text = para.text
        fullText.append(text)
    return '\n'.join(fullText)

# This is that change the .srt, .txt, .json, .md
def parse_srt(file):
    lines = file.read().decode("utf-8")
    string = lines
    text = []
    text.append(string)
    return text


def parse_epub(filename):
    book = epub.read_epub(filename)
    text = []
    for item in book.get_items():
        if item.get_type() == ebooklib.ITEM_DOCUMENT:
            bodyContent = item.get_body_content().decode()
            soup = BeautifulSoup(bodyContent, 'html.parser')
            text_ = set()  # Store unique sentences using a set
            cleaned_text = soup.get_text()
            text_.add(cleaned_text)

            text.extend(text_)
    return (text)


def text_to_docs(text, filename, chat):
    if isinstance(text, str):
        # Take a single string as one page
        text = [text]
    page_docs = [Document(page_content=page) for page in text]

    # Add page numbers as metadata
    for i, doc in enumerate(page_docs):
        doc.metadata["page"] = i + 1

    # Split pages into chunks
    doc_chunks = []

    for i, doc in enumerate(page_docs):
        text_splitter = RecursiveCharacterTextSplitter(
            separators=["\n\n", "\n"],
            chunk_size=5000,
            chunk_overlap=200,
        )
        if doc.page_content == "":
            continue
        chunks = text_splitter.split_text(doc.page_content)
        
        for i, chunk in enumerate(chunks):
            doc = Document(
                page_content=chunk
            )
            # Add sources a metadata
            doc.metadata["source"] = f"{filename}"
            doc.metadata["chat"] = f"{chat}"
            doc_chunks.append(doc)
    return doc_chunks


def web_scraping(url, max_depth=0, depth=0, visited_urls=set()):
    if depth > max_depth:  # Limit the depth to prevent infinite recursion
        return []

    try:
        # If we have already visited the url, we skip it
        if url in visited_urls:
            return []
        visited_urls.add(url)

        # Avoid unnecessary pages
        unnecessary_pages = ['contact', 'about', 'faq']
        parsed_url = urlparse(url)
        if any(page in parsed_url.path for page in unnecessary_pages):
            return []

        # Setup webdriver
        options = webdriver.ChromeOptions()
        options.add_argument("--headless")  # Run in headless mode
        options.add_argument("--no-sandbox")  # Bypass OS security model
        options.add_argument("--disable-dev-shm-usage")  # Overcome limited resource problems
        driver = webdriver.Chrome(service=ChromeService(ChromeDriverManager().install()), options=options)

        driver.get(url)
        
        # Wait for a specific element to load
        WebDriverWait(driver, 3).until(EC.presence_of_element_located((By.TAG_NAME, 'body')))
        
        soup = BeautifulSoup(driver.page_source, 'html.parser')

        text = set()  # Store unique sentences using a set
        cleaned_text = soup.get_text()
        text.add(cleaned_text)
        
        # Get all links within the page
        links = soup.find_all("a")

        # Traverse each link
        for link in links:
            href = link.get('href')
            if href is not None:
                # Create full url if href is relative
                full_url = urljoin(url, href)
                if urlparse(full_url).netloc == urlparse(url).netloc:  # Check if the link is internal
                    text.update(web_scraping(full_url, max_depth, depth + 1, visited_urls))
        
        # Generate the sentences from the set and remove unnecessary repeated text
        result = list(text)
        
        # Quit driver
        driver.quit()
        return result
    except Exception as e:
        print(f"web_scraping function error occurred: {e}")
        return []


def create_train(label, _type, status, chat, file_path=False):
    new_train = Train(label=label, type=_type, status=status, chat=chat, path=file_path)
    db.session.add(new_train)
    db.session.flush()
    db.session.commit()

    return new_train.id

def train_status_chanage(train_id):
    train_data = db.session.query(Train).filter_by(id=train_id).first()
    print(train_data, train_id)
    train_data.status = True
    db.session.commit()

def compare_role_user(chatbot):
    current_chat = db.session.query(Chat).filter_by(uuid=chatbot).first()
    traindata = json.loads(current_chat.train)
    user = db.session.query(User).filter_by(id=current_chat.user_id).first()
    # if not user.role == 1 or user.role == 7:
    #     if ct >= user.training_datas:
    #         return False
    return True

def insert_train_chat(chatbot, train_id):
    current_chat = db.session.query(Chat).filter_by(uuid=chatbot).first()
    if current_chat:
        traindata = json.loads(current_chat.train)
        traindata.append(train_id)
        current_chat.train = json.dumps(traindata)
        db.session.commit()
        chat_data = {
            'id': current_chat.id,
            'label': current_chat.label,
            'description': current_chat.description,
            'model': current_chat.model,
            'conversation': current_chat.conversation,
            'access': current_chat.access,
            'creativity': current_chat.creativity,
            'behavior': current_chat.behavior,
            'behaviormodel': current_chat.behaviormodel,
            'uuid': current_chat.uuid,
            'train': json.loads(current_chat.train),
            'chat_logo': json.loads(current_chat.chat_logo),
            'chat_title': json.loads(current_chat.chat_title),
            'chat_description': json.loads(current_chat.chat_description),
            'chat_copyright': json.loads(current_chat.chat_copyright),
            'chat_button': json.loads(current_chat.chat_button),
            'bubble': json.loads(current_chat.bubble),
            'inviteId': current_chat.inviteId,
            'api_select': current_chat.api_select
        }
        return chat_data
    else:
        return None

def duplicate_train_data(ids, chatbot_uuid):
    try:
        response = []
        for id in ids:
            train = db.session.query(Train).filter_by(id=id).first()
            if train:
                label = train.label
                train_type = train.type
                path = train.path
                new_train = Train(label=label, type=train_type, status=False, chat=chatbot_uuid, path=path)
                db.session.add(new_train)
                db.session.commit()
                
                if train_type == 'url':
                    data = web_scraping(path, 1)
                    app = current_app._get_current_object()
                    thread = Thread(target=train_data_thread, args=(new_train.id, data, path, chatbot_uuid, app))
                    thread.start()
                elif train_type == 'file':
                    file_obj = S3_CLIENT.get_object(Bucket=S3_PRIVATE_BUCKET, Key=path)
                    file_content = file_obj['Body'].read()
                    file = BytesIO(file_content)
                    if (path.split('.')[-1].lower() == 'pdf'):
                        output = parse_pdf(file)
                    elif (path.split('.')[-1].lower() == 'csv'):
                        output = parse_csv(file)
                    elif (path.split('.')[-1].lower() == 'docx'):
                        output = parse_docx(file)
                    elif path.split('.')[-1].lower() in ['srt', 'txt', 'md', 'json']:
                        output = parse_srt(file)
                    elif (path.split('.')[-1].lower() == 'epub'):
                        output = parse_epub(path.split('/')[-1])
                    app = current_app._get_current_object()
                    thread = Thread(target=train_data_thread, args=(new_train.id, output, path.split('/')[-1], chatbot_uuid, app))
                    thread.start()
                else:
                    app = current_app._get_current_object()
                    thread = Thread(target=train_data_thread, args=(new_train.id, path, path[:20], chatbot_uuid, app))
                    thread.start()
                response.append(new_train.id)
        return response
    except Exception as e:
        print(f"duplicate_train_data function error occurred: {e}")
        return []

def train_data_thread(train_id, output, filename, chatbot, app):
    with app.app_context():  # This pushes an application context
        try:
            result = text_to_docs(output, filename, chatbot)
            create_vector(result)
            train_status_chanage(train_id)
            print('Finish')
        except Exception as e:
            print(f"train_data_thread function error occurred: {e}")

@train.route('/api/data/sendurl', methods=['POST'])
def create_train_url():
    url = request.json['url']
    chatbot = request.json['chatbot']
    data = web_scraping(url, 1)
    if data == False:
        return jsonify({
            'success': False,
            'code': 405,
            'message': 'Invalid URL.',
        })

    trainid = create_train(url, 'url', False, chatbot, url)
    if (trainid == False):
        return jsonify({
            'success': False,
            'code': 405,
            'message': 'Training data already exist.',
        })
    chat = insert_train_chat(chatbot, trainid)
    app = current_app._get_current_object()
    thread = Thread(target=train_data_thread, args=(trainid, data, url, chatbot, app))
    thread.start()

    return jsonify({
        'success': True,
        'code': 200,
        'data': chat,
        'message': "create train successfully",
    })


@train.route('/api/data/sendtext', methods=['POST'])
def create_train_text():
    text = request.json['text']
    chatbot = request.json['chatbot']
    trainid = create_train(text[:20], 'text', True, chatbot, text)
    result = text_to_docs(text, text[:20], chatbot)
    if (trainid == False):
        return jsonify({
            'success': False,
            'code': 405,
            'message': 'Training data already exist.',
        })
    create_vector(result)
    chat = insert_train_chat(chatbot, trainid)
    return jsonify({
        'success': True,
        'code': 200,
        'data': chat,
        'message': "create train successfully",
    })

@train.route('/api/data/sendfile', methods=['POST'])
def create_train_file():
    try:
        filename = request.json['filename']
        chatbot = request.json['chatbot']

        file_obj = S3_CLIENT.get_object(Bucket=S3_PRIVATE_BUCKET, Key=filename)
        file_content = file_obj['Body'].read()
        file = BytesIO(file_content)
        if (filename.split('.')[-1].lower() == 'pdf'):
            output = parse_pdf(file)
        elif (filename.split('.')[-1].lower() == 'csv'):
            output = parse_csv(file)
        elif (filename.split('.')[-1].lower() == 'docx'):
            output = parse_docx(file)
        elif filename.split('.')[-1].lower() in ['srt', 'txt', 'md', 'json']:
            output = parse_srt(file)
        elif (filename.split('.')[-1].lower() == 'epub'):
            output = parse_epub(filename)

        
        trainid = create_train(filename, 'file', False, chatbot, filename)
        if (trainid == False):
            return jsonify({
                'success': False,
                'code': 401,
                'message': 'Training data already exist.',
            })
        chat = insert_train_chat(chatbot, trainid)
        app = current_app._get_current_object()
        thread = Thread(target=train_data_thread, args=(trainid, output, filename, chatbot, app))
        thread.start()
        
        return jsonify({
            'success': True,
            'code': 200,
            'data': chat,
            'message': "Train data transmission was successful. Processing in background.",
        })
    except Exception as e:
        print(f"create_train_file function error occurred: {e}")
        return {"success": False, "message": str(e)}, 400

@train.route('/api/data/sendapi', methods=['POST'])
def choose_api_for_account():
    chatbot = request.json['chatbot']  # chatbot uuid
    api = request.json['api'] # api (type: number)
    chat = db.session.query(Chat).filter_by(uuid=chatbot).first()
    chat.api_select = api
    message = "select API" if api == 1 else "disable API"
    db.session.commit()
    new_chat = {
        'id': chat.id,
        'label': chat.label,
        'description': chat.description,
        'model': chat.model,
        'conversation': chat.conversation,
        'access': chat.access,
        'creativity': chat.creativity,
        'behavior': chat.behavior,
        'behaviormodel': chat.behaviormodel,
        'uuid': chat.uuid,
        'train': json.loads(chat.train),
        'chat_logo': json.loads(chat.chat_logo),
        'chat_title': json.loads(chat.chat_title),
        'chat_description': json.loads(chat.chat_description),
        'chat_copyright': json.loads(chat.chat_copyright),
        'chat_button': json.loads(chat.chat_button),
        'bubble': json.loads(chat.bubble),
        'inviteId': chat.inviteId,
        'api_select': chat.api_select
    }
    
    return jsonify({
            'success': True,
            'code': 200,
            'data': new_chat,
            'message': message,
        })
    

@train.route('/api/data/gettraindatas', methods=['POST'])
def get_traindatas():
    uuid = request.json.get('uuid')
    if uuid:
        chat = db.session.query(Chat).filter_by(uuid=uuid).first()
        if not chat:
            return jsonify({'success': False, 'message': 'Not found', 'code': 404})
        train_ids = json.loads(chat.train)
        
        data = []
        
        user = db.session.query(User).filter_by(id=chat.user_id).first()
        invite_account = db.session.query(Invite).filter_by(email=user.email).first()
        user_check = db.session.query(User).filter_by(id=invite_account.user_id).first() if invite_account else user
        # if user_check and user_check.role == 7 and chat.api_select == 1:
        #     user = user_check
        #     if user.wonde_key:
        #         data.append({'label': 'Wonde API',
        #                     'type': 'API', 'status': True})
        for id in train_ids:
            train_data = db.session.query(Train).filter_by(id=id).first()
            if train_data:
                if train_data.type == 'API':
                    if user_check and user_check.role == 7 and chat.api_select == 1:
                        user = user_check
                        if user.wonde_key:
                            data.append({'id': train_data.id, 'label': train_data.label, 'type': train_data.type, 'status': train_data.status})
                else:
                    data.append({'id': train_data.id, 'label': train_data.label,
                        'type': train_data.type, 'status': train_data.status})
        return jsonify({'data': data, 'success': True})
    else:
        return jsonify({'success': False, 'message': 'Not found', 'code': 404})

@train.route('/api/data/deletetrain', methods=['POST'])
def delete_traindatas():
    uuid = request.json['uuid']
    id = request.json.get('id')
    chat = db.session.query(Chat).filter_by(uuid=uuid).first()
    if id:
        train_ids = json.loads(chat.train)
        train_ids.remove(id)
        chat.train = json.dumps(train_ids)
        source = db.session.query(Train).filter_by(id=id).first()
        if source.type != "API":
            # delete vectors in the pinecone
            delete_vectore(source.label, uuid)
        else:
            user = db.session.query(User).filter_by(id=chat.user_id).first()
            user.wonde_key = ''
        db.session.query(Train).filter_by(id=id).delete()
        db.session.commit()
        chat_data = {
            'id': chat.id,
            'label': chat.label,
            'description': chat.description,
            'model': chat.model,
            'conversation': chat.conversation,
            'access': chat.access,
            'creativity': chat.creativity,
            'behavior': chat.behavior,
            'behaviormodel': chat.behaviormodel,
            'uuid': chat.uuid,
            'train': json.loads(chat.train),
            'chat_logo': json.loads(chat.chat_logo),
            'chat_title': json.loads(chat.chat_title),
            'chat_description': json.loads(chat.chat_description),
            'chat_copyright': json.loads(chat.chat_copyright),
            'chat_button': json.loads(chat.chat_button),
            'bubble': json.loads(chat.bubble),
        }
        data = {
            'code': 200,
            'message': "Succesfullu delete",
            'data': chat_data,
            'success': True
        }
    else:
        chat.api_select = 0 # Diselect the wonde API key
        db.session.commit()
        chat_data = {
            'id': chat.id,
            'label': chat.label,
            'description': chat.description,
            'model': chat.model,
            'conversation': chat.conversation,
            'access': chat.access,
            'creativity': chat.creativity,
            'behavior': chat.behavior,
            'behaviormodel': chat.behaviormodel,
            'uuid': chat.uuid,
            'train': json.loads(chat.train),
            'chat_logo': json.loads(chat.chat_logo),
            'chat_title': json.loads(chat.chat_title),
            'chat_description': json.loads(chat.chat_description),
            'chat_copyright': json.loads(chat.chat_copyright),
            'chat_button': json.loads(chat.chat_button),
            'bubble': json.loads(chat.bubble),
        }
        data = {
            'code': 200,
            'message': "Succesfullu delete",
            'data': chat_data,
            'success': True
        }
    return jsonify(data)


@train.route('/api/data/deletetrain_vectore', methods=['POST'])
def delete_pinecone_vectore():

    password = request.json['password']
    source = request.json['source']
    if password == 'QWE@#$asd234':
        delete_vectore(source)
        return jsonify({'data': 'okay'})
    return jsonify({'data': 'You are unautherize'})
